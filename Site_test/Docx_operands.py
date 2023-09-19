from asyncio.windows_events import NULL
from datetime import datetime
from docx import *
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Docx_doc:

    def __init__(self, left = 20, right = 20, top = 20, bottom = 20):
       self.doc = Document()

       section = self.doc.sections[0]
       section.left_margin = Mm(left)
       section.right_margin = Mm(right)
       section.top_margin = Mm(top)
       section.bottom_margin = Mm(bottom)

       self.New_paragraph()

    def New_paragraph(self):
        self.par = self.doc.add_paragraph()
        self.run = self.par.add_run()

        paragraph_format = self.par.paragraph_format
       # paragraph_format.line_spacing = Pt(20)
        paragraph_format.space_after = Pt(1)
   
    def New_lines(self, number = 1):
        for i in range(0, number, 1):
            self.New_paragraph()
 
    def Add_many_lines(self, lines_ , newline_ = False, align_ = 'L', size_ = 12, font_ = "Times New Roman", bold_ = False, color_ = "black"):
        for line_ in lines_:
            self.Add_text(line_ , newline_, align_, size_, font_, bold_, color_)

    def Add_text(self, text_ , newline_ = False, align_ = 'L', size_ = 12, font_ = "Times New Roman", bold_ = False, color_ = "black"):

        if align_ == 'L':
            self.par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align_ == 'R':
            self.par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align_ == 'C':
            self.par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.run.font.name = font_
        self.run.font.size = Pt(size_)
        self.run.bold = bold_
        
        match (color_):
            case "black":
                font_color = RGBColor(0, 0, 0)
            case "red":
                font_color = RGBColor(255, 0, 0)
            case "white":
                font_color = RGBColor(255, 255, 255)
            case "blue":
                font_color = RGBColor(50, 150, 255)
            case "yellow":
                font_color = RGBColor(255, 255, 100)
            case "orange":
                font_color = RGBColor(255, 150, 50)
            case "violet":
                font_color = RGBColor(180, 0, 180)
                
        self.run.font.color.rgb = font_color
        self.run.add_text(text_)

        if newline_ == True:
            self.New_paragraph()

    def Add_pic(self, link, width_ = NULL):
        if width_ == NULL:
            width_ = self.doc.sections[0].page_width - self.doc.sections[0].left_margin -  \
                         self.doc.sections[0].right_margin
        else:
            width_ = Mm(width_)

        self.run.add_picture(link, width_)

    def Create_table(self, rows_, columns_, style_ = 'Table Grid'):
        self.table = self.doc.add_table(rows=rows_, cols=columns_)
        self.table.style = style_

    def Fill_cell(self, cell_, text_, font_, size_):
        cell_.text = str(text_)
        cell_.paragraphs[0].runs[0].font.name = font_
        cell_.paragraphs[0].runs[0].font.size = Pt(size_)
        
    def Fill_table(self, titles_ , rows_, cols_width_ = NULL, bold_ = True, font_ = "Times New Roman", size_ = 12, index_ = False):
        if len(rows_) == 0:
            self.Add_text("Нет объектов", True)
            return

        if index_ == True:
            cols_num = len(rows_[0]) + 1
        else:
            cols_num = len(rows_[0])

        self.table = self.doc.add_table(1, cols_num)
     
        if cols_width_ != NULL:
            width_ = self.doc.sections[0].page_width - self.doc.sections[0].left_margin -  \
                             self.doc.sections[0].right_margin
            
            sum_cols = sum(cols_width_)
            for i in range(0, len(rows_[0])):
                 self.table.columns[i].width = int(width_ * float(cols_width_[i]) / float(sum_cols))

        self.table.style = 'Table Grid'

        head_cells = self.table.rows[0].cells
        for i, title in enumerate(titles_):
            p = head_cells[i].paragraphs[0]
            p.add_run(title).bold = bold_
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.name = font_
            p.runs[0].font.size = Pt(size_)

        i = 1
        for row in rows_:
            cells = self.table.add_row().cells
            if index_ == True:
                i2 = 1
                self.Fill_cell(cells[0], str(i), font_, size_)
            else:
                i2 = 0

            for cell_data in row:
                self.Fill_cell(cells[i2], cell_data, font_, size_)
                i2 += 1
            i += 1

        self.New_paragraph()

   
    def Save_docx(self, name = NULL):
        if name == NULL:
            name = r"Reports\Report " + str(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            name = name.replace("-", ".")
            name = name.replace(":", "_")
         
        self.doc.save(name + ".docx")




            


